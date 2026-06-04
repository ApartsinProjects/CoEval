"""Finalize a two-column DOCX into TACL submission geometry.

Applies, to EVERY section (the two-column build creates ~23 section breaks for
full-width floats, so per-section is the only robust way): A4 page size, 2.5 cm
margins, margin line numbers (review requirement), and the confidentiality header.

These belong in the html2doc 'tacl' profile long-term (the page_size / line_numbers /
header_text keys are already declared there as hooks); kept here as a project finalizer
until folded into the skill's apply_academic_style post-process.

Usage:  python docs/paper/finalize_tacl.py docs/paper/CoEval_TACL.docx
"""
import sys
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HEADER = "Confidential TACL submission. DO NOT DISTRIBUTE."
path = sys.argv[1]
doc = Document(path)

for s in doc.sections:
    # A4 + 2.5 cm margins
    s.page_width, s.page_height = Mm(210), Mm(297)
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Mm(25)
    sectPr = s._sectPr
    # margin line numbers (continuous), inserted in schema order (before w:cols)
    for ex in sectPr.findall(qn("w:lnNumType")):
        sectPr.remove(ex)
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "continuous")
    # Small distance (~0.1") so the RIGHT column's numbers fit inside the 0.6cm gutter
    # without crossing into the left column. (284 twips overflowed the gutter.)
    ln.set(qn("w:distance"), "80")
    cols = sectPr.find(qn("w:cols"))
    if cols is not None:
        cols.addprevious(ln)
    else:
        sectPr.append(ln)
    # confidentiality header (unlinked so it shows on every section)
    s.header.is_linked_to_previous = False
    hp = s.header.paragraphs[0] if s.header.paragraphs else s.header.add_paragraph()
    hp.text = HEADER
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in hp.runs:
        r.font.size = Pt(9)
        r.font.name = "Times New Roman"
        r.font.italic = True

# Make line numbers small (7pt) so the right-column numbers fit the narrow gutter
# without overlapping the left column. Word formats line numbers via the BUILT-IN
# character style with styleId "LineNumber" -- python-docx's add_style('Line Number')
# makes a phantom style Word ignores, so inject the real styleId directly via OXML.
styles_el = doc.styles.element
_ln = None
for _st in styles_el.findall(qn("w:style")):
    if _st.get(qn("w:styleId")) == "LineNumber":
        _ln = _st
        break
if _ln is None:
    _ln = OxmlElement("w:style")
    _ln.set(qn("w:type"), "character")
    _ln.set(qn("w:styleId"), "LineNumber")
    _nm = OxmlElement("w:name")
    _nm.set(qn("w:val"), "Line Number")
    _ln.append(_nm)
    styles_el.append(_ln)
_rPr = _ln.find(qn("w:rPr"))
if _rPr is None:
    _rPr = OxmlElement("w:rPr")
    _ln.append(_rPr)
for _tag, _val in (("w:rFonts", None), ("w:sz", "14"), ("w:szCs", "14")):
    _e = _rPr.find(qn(_tag))
    if _e is None:
        _e = OxmlElement(_tag)
        _rPr.append(_e)
    if _tag == "w:rFonts":
        _e.set(qn("w:ascii"), "Times New Roman")
        _e.set(qn("w:hAnsi"), "Times New Roman")
    else:
        _e.set(qn("w:val"), _val)  # 14 half-points = 7pt

# Compact the figures to reclaim vertical space: cap each chart's height at 2.2 in
# (full-width body charts render ~3.3 in tall, eating page space). Keeps aspect ratio;
# charts stay legible. Appendix figures shrink too (harmless).
MAX_H_IN = 2.0
for sh in doc.inline_shapes:
    h_in = sh.height / 914400.0
    if h_in > MAX_H_IN:
        scale = MAX_H_IN / h_in
        sh.height = int(sh.height * scale)
        sh.width = int(sh.width * scale)

doc.save(path)

# report
doc2 = Document(path)
import zipfile
x = zipfile.ZipFile(path).read("word/document.xml").decode("utf-8")
print("sections:", len(doc2.sections))
print("A4 width (11906 twips):", all(round(s.page_width.twips) == 11906 for s in doc2.sections))
print("line numbers on all sections:", x.count("lnNumType"), "of", len(doc2.sections))
print("header set:", HEADER in zipfile.ZipFile(path).read(
    [n for n in zipfile.ZipFile(path).namelist() if "header" in n][0]).decode("utf-8")
    if any("header" in n for n in zipfile.ZipFile(path).namelist()) else False)
