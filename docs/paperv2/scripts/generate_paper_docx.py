#!/usr/bin/env python3
"""
generate_paper_docx.py
======================
Generates a single-column Microsoft Word document for:

  CoEval: A Self-Evaluating LLM Ensemble Framework for
  Scalable, Attribute-Controlled Benchmark Generation

Formatting follows ACL 2026 Long Paper submission guidelines:
  - Letter paper (8.5" x 11"), 1" margins
  - Times New Roman, 10pt body / 12pt section headers / 16pt title
  - Single-column layout throughout
  - Figures with high-quality captions and in-text cross-references
  - Footnotes for all simulated results (marked with *)
  - Disclaimer section at end of paper

Usage:
    python generate_paper_docx.py [--output PATH]
"""

import re
import sys
import copy
import argparse
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# ──────────────────────────────────────────────────────────────────────────────
# Paths
# ──────────────────────────────────────────────────────────────────────────────
PAPER_DIR  = Path(__file__).resolve().parent.parent
FIGS_DIR   = PAPER_DIR / "figures"
SCRIPT_DIR = Path(__file__).resolve().parent

SECTION_FILES = {
    "abstract":     PAPER_DIR / "01_abstract.md",
    "introduction": PAPER_DIR / "02_introduction.md",
    "related_work": PAPER_DIR / "03_related_work_v2.md",   # v2: ARES, Verga, FLAMe fixed
    "methodology":  PAPER_DIR / "04_methodology_v2.md",    # v2: Eqs 1-10, WPA fixed, D* fixed
    "experiments":  PAPER_DIR / "05_experiments_v2.md",    # v2: §4.6 demoted, ablation table
    "analysis":     PAPER_DIR / "06_analysis_conclusion_v2.md",  # v2: academic voice, kappa anchor
    "appendix":     PAPER_DIR / "10_appendix.md",
    "disclaimer":   PAPER_DIR / "10_disclaimer.md",         # moved from hardcoded constant
    "references":   PAPER_DIR / "references.md",
}

# ──────────────────────────────────────────────────────────────────────────────
# Paper metadata loader
# ──────────────────────────────────────────────────────────────────────────────

def _parse_paper_metadata(filepath: Path) -> dict:
    """
    Parse 00_paper_metadata.md into a dict of SECTION_KEY → stripped text.
    Sections are delimited by '## ALL_CAPS_KEY' lines.
    Multi-line values are joined with '\\n'.
    """
    result = {}
    try:
        lines = filepath.read_text(encoding="utf-8").splitlines()
    except FileNotFoundError:
        print(f"  [WARN] Metadata file not found: {filepath}")
        return result
    _key_re = re.compile(r'^## ([A-Z][A-Z0-9_]+)$')
    current_key = None
    current_lines = []
    for line in lines:
        m = _key_re.match(line.strip())
        if m:
            if current_key:
                result[current_key] = "\n".join(current_lines).strip()
            current_key = m.group(1)
            current_lines = []
        elif current_key and not line.startswith("#") and not line.startswith("<!--"):
            current_lines.append(line)
    if current_key:
        result[current_key] = "\n".join(current_lines).strip()
    return result


_META = _parse_paper_metadata(PAPER_DIR / "00_paper_metadata.md")

REPO_URL          = _META.get("REPO_URL",       "https://github.com/ApartsinProjects/CoEval")

# ──────────────────────────────────────────────────────────────────────────────
# Figure catalogue — maps label → (path, caption, in_text_ref)
# Each figure: (file_path, caption, width_inches)
# ──────────────────────────────────────────────────────────────────────────────
FIGURES = {
    "fig1": (
        FIGS_DIR / "diagrams" / "architecture.png",
        "Figure 1: CoEval pipeline architecture. Teachers generate attribute-controlled "
        "benchmark items; students provide responses; judges score independently against "
        "rubric criteria. Role assignments rotate across ensemble members.",
        5.5,
    ),
    "fig2_tsj": (
        FIGS_DIR / "diagrams" / "fig_teacher_student_judge.png",
        "Figure 2: Teacher–Student–Judge role separation. Any LLM may participate in "
        "multiple roles simultaneously within a single pipeline run.",
        4.5,
    ),
    "fig3_overview": (
        FIGS_DIR / "screenshots" / "fig_overview.png",
        "Figure 3: Example CoEval overview report (medium-benchmark-v1). Displays "
        "aggregate statistics across all tasks, models, and evaluation phases.",
        6.5,
    ),
    "fig4_judge_agree": (
        FIGS_DIR / "screenshots" / "fig_judge_agreement_top.png",
        "Figure 4: Example judge agreement report (medium-benchmark-v1). Top judge "
        "pairs by Weighted Pairwise Agreement (WPA). GPT-3.5-Turbo × GPT-4o-mini "
        "achieves the highest WPA (0.812) and Cohen's κ (0.422).",
        6.5,
    ),
    "fig5_judge_consistency": (
        FIGS_DIR / "screenshots" / "fig_judge_consistency.png",
        "Figure 5 (Appendix): Example per-judge consistency report (medium-benchmark-v1). "
        "SmolLM2-1.7B exhibits near-random pairwise agreement (κ ≈ 0.003) while "
        "GPT-4o-mini shows substantially higher reliability.",
        6.5,
    ),
    "fig6_score_dist": (
        FIGS_DIR / "screenshots" / "fig_score_distribution.png",
        "Figure 6 (Appendix): Example score distribution report (medium-benchmark-v1). "
        "Data interpretation consistently yields the lowest scores (0.63–0.73); "
        "text summarization yields the highest (0.79–0.87).",
        6.5,
    ),
    "fig7_teacher": (
        FIGS_DIR / "screenshots" / "fig_teacher_report.png",
        "Figure 7 (Appendix): Example teacher discrimination report (medium-benchmark-v1). "
        "SmolLM2-1.7B achieves the highest V1 score (0.0046), indicating prompt "
        "diversity rather than generation quality drives discrimination power.",
        6.5,
    ),
    "fig8_ensemble_abl": (
        FIGS_DIR / "tables" / "fig_ensemble_ablation_simulated.png",
        "Figure 8: J* filter and OLS calibration ablation — projected values (*). "
        "Results are placeholders based on design targets; will be replaced with "
        "measured experimental results.",
        5.5,
    ),
    "fig9_benchmark_comp": (
        FIGS_DIR / "tables" / "fig_benchmark_comparison_simulated.png",
        "Figure 9: Comparative evaluation against G-Eval and BERTScore — projected "
        "values (*). Comparison against held-out human annotations. Results are "
        "placeholders; will be replaced with measured experimental results.",
        5.5,
    ),
    "fig10_pos_bias": (
        FIGS_DIR / "tables" / "fig_positional_bias_simulated.png",
        "Figure 10: Positional bias analysis — projected values (*). OLS calibration "
        "is projected to reduce positional flip rate from ≈23% to under 5%. "
        "Results are placeholders; will be replaced with measured experimental results.",
        5.5,
    ),
    "fig11_kappa_matrix": (
        FIGS_DIR / "tables" / "fig_judge_agreement_matrix.png",
        "Figure 11: Inter-judge Cohen's κ matrix (medium-benchmark-v1). Large-model "
        "pairs (GPT family) consistently achieve κ > 0.3; cross-scale pairs show "
        "near-zero agreement.",
        5.5,
    ),
    "fig12_teacher_disc": (
        FIGS_DIR / "tables" / "fig_teacher_discrimination.png",
        "Figure 12: Teacher discrimination scores V1, S2, R3 (medium-benchmark-v1). "
        "SmolLM2-1.7B ranks first on all three metrics despite producing lower-quality "
        "individual items — a counter-intuitive finding.",
        5.5,
    ),
    "fig13_cost": (
        FIGS_DIR / "tables" / "fig_cost_breakdown.png",
        "Figure 13: API cost breakdown by phase and provider (medium-benchmark-v1). "
        "The evaluation phase dominates total cost ($4.48 of $5.89 total).",
        5.5,
    ),
    "fig14_student": (
        FIGS_DIR / "screenshots" / "fig_student_report.png",
        "Figure 14 (Appendix): Example student performance report (medium-benchmark-v1). "
        "GPT-4o-mini achieves the highest mean score (0.832); Qwen2.5-0.5B the "
        "lowest (0.648).",
        6.5,
    ),
    "fig15_coverage": (
        FIGS_DIR / "screenshots" / "fig_coverage_summary.png",
        "Figure 15 (Appendix): Example attribute coverage report (medium-benchmark-v1). "
        "All target attribute combinations are covered, validating CoEval's stratified "
        "sampling mechanism.",
        6.5,
    ),
    # ── Appendix figures: wizard console + detailed report screenshots ──────────
    "fig16_wizard1": (
        FIGS_DIR / "screenshots" / "fig_wizard_step1.png",
        "Figure 16 (Appendix): CoEval wizard — task description and attribute axis "
        "configuration. The wizard elicits a task description, infers target and "
        "nuanced attributes, and builds the YAML configuration interactively.",
        6.5,
    ),
    "fig17_wizard2": (
        FIGS_DIR / "screenshots" / "fig_wizard_step2.png",
        "Figure 17 (Appendix): CoEval wizard — model selection and pipeline "
        "configuration. Teachers, students, and judges are assigned from available "
        "provider pools before generating the final YAML.",
        6.5,
    ),
    "fig18_yaml_detail": (
        FIGS_DIR / "screenshots" / "fig_yaml_config_detailed.png",
        "Figure 18 (Appendix): Example YAML configuration for the medium-benchmark-v1 "
        "experiment. The declarative format specifies task definitions, attribute axes, "
        "model assignments, and sampling parameters in a single reproducible artifact.",
        6.5,
    ),
    "fig19_judge_detail": (
        FIGS_DIR / "screenshots" / "fig_judge_report_detail.png",
        "Figure 19 (Appendix): Example detailed judge consistency report "
        "(medium-benchmark-v1). Shows SPA, WPA, and κ per judge and per criterion. "
        "'Professionalism' exhibits the lowest cross-judge agreement (SPA = 0.294).",
        6.5,
    ),
    "fig20_interaction_detail": (
        FIGS_DIR / "screenshots" / "fig_interaction_matrix_detail.png",
        "Figure 20 (Appendix): Example judge–criterion interaction matrix "
        "(medium-benchmark-v1). Heatmap cells show mean WPA for each (judge, criterion) "
        "combination.",
        6.5,
    ),
    "fig21_coverage_detail": (
        FIGS_DIR / "screenshots" / "fig_coverage_detail.png",
        "Figure 21 (Appendix): Example detailed attribute coverage report "
        "(medium-benchmark-v1). Per-cell counts for all target and nuanced attribute "
        "combinations.",
        6.5,
    ),
    "fig22_summary_detail": (
        FIGS_DIR / "screenshots" / "fig_summary_detail.png",
        "Figure 22 (Appendix): Example full summary report (medium-benchmark-v1). "
        "Shows phase completion status, item counts, total cost, and per-model "
        "aggregate scores.",
        6.5,
    ),
}

# ──────────────────────────────────────────────────────────────────────────────
# Simulated-result footnote marker and text
# ──────────────────────────────────────────────────────────────────────────────
SIMULATED_FN_TEXT = _META.get(
    "SIMULATED_FN_TEXT",
    "Results marked with (*) are simulated — see Disclaimer section for details."
)

# DISCLAIMER_TEXT has been moved to 10_disclaimer.md
# Edit that file to update the disclaimer.


# ──────────────────────────────────────────────────────────────────────────────
# XML helpers
# ──────────────────────────────────────────────────────────────────────────────

def _make_element(tag, **attrib):
    el = OxmlElement(tag)
    for k, v in attrib.items():
        el.set(qn(k), v)
    return el


def _set_cols(section, num_cols=2, space_twips=720):
    """Set the number of text columns in a Word section."""
    sectPr = section._sectPr
    # Remove any existing cols element
    for old in sectPr.findall(qn("w:cols")):
        sectPr.remove(old)
    cols = _make_element("w:cols", **{"w:num": str(num_cols),
                                      "w:space": str(space_twips),
                                      "w:equalWidth": "1"})
    sectPr.append(cols)


def _set_margins(section, top=1.0, bottom=1.0, left=1.0, right=1.0):
    """Set page margins (in inches)."""
    sectPr = section._sectPr
    pgMar = sectPr.find(qn("w:pgMar"))
    if pgMar is None:
        pgMar = OxmlElement("w:pgMar")
        sectPr.append(pgMar)
    for attr, val in [("w:top",    int(top * 1440)),
                      ("w:bottom", int(bottom * 1440)),
                      ("w:left",   int(left * 1440)),
                      ("w:right",  int(right * 1440)),
                      ("w:header", 720),
                      ("w:footer", 720)]:
        pgMar.set(qn(attr), str(val))


def _add_page_numbers(section):
    """Add centered page number to the footer of a section."""
    footer = section.footer
    if not footer.paragraphs:
        footer.add_paragraph()
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    _apply_font(run, size_pt=9)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_continuous_section_break(doc):
    """Insert a continuous section break at the current position."""
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    sPr  = OxmlElement("w:sectPr")
    cols = _make_element("w:cols", **{"w:num": "2", "w:space": "360"})
    sPr.append(cols)
    pgSz = _make_element("w:pgSz", **{"w:w": "12240", "w:h": "15840"})
    sPr.append(pgSz)
    pgMar = _make_element("w:pgMar", **{
        "w:top": "1440", "w:right": "1440",
        "w:bottom": "1440", "w:left": "1440",
        "w:header": "720", "w:footer": "720",
    })
    sPr.append(pgMar)
    pPr.append(sPr)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    return para


# ── Footnote tracker (simple numbered superscript approach, universally reliable)
class _FootnoteTracker:
    """
    Tracks footnotes as a list of (text, mark) tuples.
    Repeated identical texts reuse the same mark.
    Marks for simulated results always use '*'.
    """
    _SIMULATED_MARK = "*"

    def __init__(self):
        self._registry = {}   # text -> mark string
        self._ordered  = []   # [(mark, text), ...]
        self._counter  = 0

    def add(self, para, text, is_simulated=False):
        """Add a superscript footnote mark to `para` and register the text."""
        if text in self._registry:
            mark = self._registry[text]
        elif is_simulated or text == SIMULATED_FN_TEXT:
            mark = self._SIMULATED_MARK
            if text not in self._registry:
                self._registry[text] = mark
                self._ordered.append((mark, text))
        else:
            self._counter += 1
            mark = str(self._counter)
            self._registry[text] = mark
            self._ordered.append((mark, text))

        # Add superscript mark to paragraph
        run = para.add_run(f"[{mark}]")
        rpr = run._r.get_or_add_rPr()
        va  = OxmlElement("w:vertAlign")
        va.set(qn("w:val"), "superscript")
        rpr.append(va)
        _apply_font(run, size_pt=7)

    def write_to_doc(self, doc):
        """Write all footnotes as a formatted block at current position."""
        if not self._ordered:
            return
        seen = set()
        for mark, text in self._ordered:
            if mark in seen:
                continue
            seen.add(mark)
            p = doc.add_paragraph()
            _para_fmt(p, space_before=0, space_after=2)
            r1 = p.add_run(f"[{mark}] ")
            _apply_font(r1, size_pt=8, bold=True)
            r2 = p.add_run(text)
            _apply_font(r2, size_pt=8, italic=False)

_FN = _FootnoteTracker()   # global tracker for the document


def _add_footnote(doc, para, footnote_text, is_simulated=False):
    """Add a superscript footnote reference to `para`."""
    _FN.add(para, footnote_text, is_simulated=is_simulated)


# ──────────────────────────────────────────────────────────────────────────────
# Style helpers
# ──────────────────────────────────────────────────────────────────────────────

def _apply_font(run, name="Times New Roman", size_pt=10, bold=False, italic=False,
                color=None):
    run.font.name     = name
    run.font.size     = Pt(size_pt)
    run.font.bold     = bold
    run.font.italic   = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Ensure East-Asian fonts also use the same face
    rpr = run._r.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"),   name)
    rFonts.set(qn("w:hAnsi"),   name)
    rFonts.set(qn("w:eastAsia"), name)


def _set_font_meta(run, name="Times New Roman", size_pt=10):
    """Set font name and size ONLY — preserves existing bold/italic state.
    Use this after _render_inline to avoid wiping inline formatting."""
    run.font.name = name
    run.font.size = Pt(size_pt)
    rpr = run._r.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"),   name)
    rFonts.set(qn("w:hAnsi"),   name)
    rFonts.set(qn("w:eastAsia"), name)


def _para_fmt(para, space_before=0, space_after=6, line_spacing=1.0,
              alignment=WD_ALIGN_PARAGRAPH.LEFT, keep_with_next=False):
    pf = para.paragraph_format
    pf.space_before    = Pt(space_before)
    pf.space_after     = Pt(space_after)
    pf.alignment       = alignment
    pf.keep_with_next  = keep_with_next
    if line_spacing == 1.0:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    else:
        pf.line_spacing = Pt(12 * line_spacing)


def _add_heading(doc, text, level=1, numbered=True, number_prefix=""):
    """Add a section heading with Times New Roman formatting."""
    style_name = f"Heading {level}"
    try:
        para = doc.add_paragraph(style=style_name)
    except Exception:
        para = doc.add_paragraph()
    para.clear()
    run = para.add_run(text)
    sizes = {1: 12, 2: 11, 3: 10}
    space_befores = {1: 10, 2: 6, 3: 4}
    space_afters  = {1: 4,  2: 3, 3: 2}
    _apply_font(run, size_pt=sizes.get(level, 10), bold=True)
    _para_fmt(para, space_before=space_befores.get(level, 6),
              space_after=space_afters.get(level, 3),
              keep_with_next=True,
              alignment=WD_ALIGN_PARAGRAPH.LEFT)
    return para


def _add_normal(doc, text, bold_spans=None, italic_spans=None,
                first_line_indent=False):
    """Add a normal body paragraph, optionally with inline bold/italic spans."""
    para = doc.add_paragraph()
    _para_fmt(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    if first_line_indent:
        para.paragraph_format.first_line_indent = Pt(12)
    _render_inline(para, text)
    return para


def _preprocess_latex(text):
    """
    Convert LaTeX markup to plain-text equivalents for Word rendering.
    Handles citations, fractions, binomials, Greek letters, accented symbols,
    and math operators.  Accented commands (hat, bar, tilde) are applied AFTER
    Greek conversion so that e.g. \\hat{\\sigma} correctly yields σ̂.
    """
    # ── Citations ────────────────────────────────────────────────────────
    text = re.sub(r'\\citep\{([^}]+)\}', r'[\1]', text)
    text = re.sub(r'\\citet\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\citeauthor\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\citealp\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\cite\{([^}]+)\}', r'[\1]', text)
    # ── Equation tags ────────────────────────────────────────────────────
    text = re.sub(r'\\tag\{(\d+)\}', r'  (Eq. \1)', text)
    # ── Binomial coefficients ─────────────────────────────────────────────
    for _ in range(3):
        text = re.sub(r'\\binom\{([^{}]*)\}\{([^{}]*)\}', r'C(\1,\2)', text)
    # ── Fractions ────────────────────────────────────────────────────────
    for _ in range(4):
        text = re.sub(r'\\frac\{([^{}]*)\}\{([^{}]*)\}', r'(\1/\2)', text)
    # ── Named math operators ─────────────────────────────────────────────
    text = re.sub(r'\\operatorname\*?\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\text\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\mathcal\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\mathbf\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\mathrm\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\boldsymbol\{([^}]+)\}', r'\1', text)
    # ── Blackboard bold (ℕ, ℝ, etc.) ────────────────────────────────────
    _BB = {'N': 'ℕ', 'Z': 'ℤ', 'Q': 'ℚ', 'R': 'ℝ', 'C': 'ℂ', 'E': '𝔼', 'P': 'ℙ'}
    text = re.sub(r'\\mathbb\{([A-Z])\}', lambda m: _BB.get(m.group(1), m.group(1)), text)
    text = re.sub(r'\\mathbb\{([^}]+)\}', r'\1', text)
    # ── Subscripts / superscripts ─────────────────────────────────────────
    text = re.sub(r'_\{([^{}]*)\}', r'_\1', text)
    text = re.sub(r'\^\{([^{}]*)\}', r'^\1', text)
    # ── Delimiters ──────────────────────────────────────────────────────
    text = text.replace(r'\left(', '(')
    text = text.replace(r'\right)', ')')
    text = text.replace(r'\left[', '[')
    text = text.replace(r'\right]', ']')
    text = text.replace(r'\left\{', '{')
    text = text.replace(r'\right\}', '}')
    text = text.replace(r'\left|', '|')
    text = text.replace(r'\right|', '|')
    text = text.replace(r'\lceil', '⌈')
    text = text.replace(r'\rceil', '⌉')
    text = text.replace(r'\lfloor', '⌊')
    text = text.replace(r'\rfloor', '⌋')
    for tok in (r'\bigl', r'\bigr', r'\Bigl', r'\Bigr', r'\big', r'\Big',
                r'\bigg', r'\Bigg', r'\!\!', r'\!', r'\,', r'\;', r'\:',
                r'\quad', r'\qquad'):
        text = text.replace(tok, ' ')
    # ── Greek letters ────────────────────────────────────────────────────
    greeks = {
        r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\delta': 'δ',
        r'\epsilon': 'ε', r'\varepsilon': 'ε', r'\zeta': 'ζ', r'\eta': 'η',
        r'\theta': 'θ', r'\vartheta': 'θ', r'\iota': 'ι', r'\kappa': 'κ',
        r'\lambda': 'λ', r'\mu': 'μ', r'\nu': 'ν', r'\xi': 'ξ',
        r'\pi': 'π', r'\rho': 'ρ', r'\sigma': 'σ', r'\tau': 'τ',
        r'\upsilon': 'υ', r'\phi': 'φ', r'\varphi': 'φ', r'\chi': 'χ',
        r'\psi': 'ψ', r'\omega': 'ω',
        r'\Gamma': 'Γ', r'\Delta': 'Δ', r'\Theta': 'Θ', r'\Lambda': 'Λ',
        r'\Xi': 'Ξ', r'\Pi': 'Π', r'\Sigma': 'Σ', r'\Upsilon': 'Υ',
        r'\Phi': 'Φ', r'\Psi': 'Ψ', r'\Omega': 'Ω',
    }
    for cmd, sym in greeks.items():
        text = text.replace(cmd, sym)
    # ── Math symbols ─────────────────────────────────────────────────────
    symbols = {
        r'\approx': '≈', r'\leq': '≤', r'\geq': '≥', r'\neq': '≠',
        r'\le': '≤', r'\ge': '≥', r'\ne': '≠', r'\ll': '≪', r'\gg': '≫',
        r'\infty': '∞', r'\times': '×', r'\cdot': '·', r'\circ': '∘',
        r'\in': '∈', r'\notin': '∉', r'\subset': '⊂', r'\subseteq': '⊆',
        r'\supset': '⊃', r'\supseteq': '⊇', r'\cup': '∪', r'\cap': '∩',
        r'\sum': 'Σ', r'\prod': 'Π', r'\int': '∫', r'\partial': '∂',
        r'\nabla': '∇', r'\forall': '∀', r'\exists': '∃',
        r'\pm': '±', r'\mp': '∓', r'\div': '÷',
        r'\mapsto': '↦', r'\hookrightarrow': '↪',
        r'\rightarrow': '→', r'\leftarrow': '←', r'\leftrightarrow': '↔',
        r'\Rightarrow': '⇒', r'\Leftarrow': '⇐', r'\Leftrightarrow': '⇔',
        r'\to': '→', r'\gets': '←',
        r'\ldots': '…', r'\cdots': '…', r'\vdots': '⋮', r'\ddots': '⋱',
        r'\langle': '⟨', r'\rangle': '⟩',
        r'\|': '‖', r'\ell': 'ℓ', r'\Re': 'ℜ', r'\Im': 'ℑ',
        r'\top': '⊤', r'\bot': '⊥', r'\emptyset': '∅', r'\varnothing': '∅',
        r'\neg': '¬', r'\land': '∧', r'\lor': '∨',
        r'\sqrt': '√', r'\perp': '⊥',
    }
    for cmd, sym in symbols.items():
        text = text.replace(cmd, sym)
    # ── Accented symbols — processed AFTER Greek so σ, μ etc. already converted ──
    def _hat_repl(m):
        c = m.group(1).strip()
        return (c + '\u0302') if len(c) == 1 else f'hat({c})'
    def _bar_repl(m):
        c = m.group(1).strip()
        return (c + '\u0304') if len(c) == 1 else f'mean({c})'
    def _tilde_repl(m):
        c = m.group(1).strip()
        return (c + '\u0303') if len(c) == 1 else f'~{c}'
    for _ in range(3):
        text = re.sub(r'\\hat\{([^{}]*)\}', _hat_repl, text)
        text = re.sub(r'\\widehat\{([^{}]*)\}', _hat_repl, text)
        text = re.sub(r'\\bar\{([^{}]*)\}', _bar_repl, text)
        text = re.sub(r'\\overline\{([^{}]*)\}', _bar_repl, text)
        text = re.sub(r'\\tilde\{([^{}]*)\}', _tilde_repl, text)
        text = re.sub(r'\\vec\{([^{}]*)\}', r'\1', text)
    # ── Clean up remaining braces and backslash commands ─────────────────
    for _ in range(4):
        text = re.sub(r'\{([^{}]*)\}', r'\1', text)
    text = re.sub(r'\\[a-zA-Z]+\*?', '', text)  # remaining unknown commands
    # ── Inline math delimiters ───────────────────────────────────────────
    text = re.sub(r'\$([^\$\n]+)\$', r'\1', text)
    # ── LaTeX tilde non-breaking space → regular space ───────────────────
    text = re.sub(r'(?<!\\)~', ' ', text)
    # ── Normalize whitespace ─────────────────────────────────────────────
    text = re.sub(r'  +', ' ', text).strip()
    return text


def _render_inline(para, text):
    """
    Parse simple inline markdown (**bold**, *italic*, `code`)
    and add runs to `para`. Preprocesses LaTeX citations and math macros.
    """
    text = _preprocess_latex(text)
    # Pattern: **bold**, *italic*, `code`, or plain text
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
    parts   = re.split(pattern, text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            _apply_font(run, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            _apply_font(run, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            _apply_font(run, name="Courier New", size_pt=9)
        else:
            if part:
                run = para.add_run(part)
                _apply_font(run)


def _add_caption(doc, text, fig_label=None):
    """Add a figure caption paragraph."""
    para = doc.add_paragraph()
    _para_fmt(para, space_before=2, space_after=4,
              alignment=WD_ALIGN_PARAGRAPH.CENTER)
    para.paragraph_format.left_indent  = Inches(0)
    para.paragraph_format.right_indent = Inches(0)
    run = para.add_run(text)
    _apply_font(run, size_pt=9, italic=True)
    # Add simulated footnote if caption contains * (simulated placeholder)
    if "projected values (*)" in text.lower() or "placeholders" in text.lower() or \
       (text.count("(*)") > 0) or text.endswith("*"):
        _add_footnote(doc, para, SIMULATED_FN_TEXT)
    return para


def _add_figure(doc, fig_key):
    """Insert a figure with caption. Returns the caption paragraph."""
    if fig_key not in FIGURES:
        print(f"  [WARN] Figure key '{fig_key}' not in catalogue — skipped.")
        return None
    fig_path, caption, width_in = FIGURES[fig_key]
    if not Path(fig_path).exists():
        print(f"  [WARN] Figure file not found: {fig_path}")
        # Add placeholder paragraph
        p = doc.add_paragraph()
        _para_fmt(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        p.paragraph_format.left_indent  = Inches(0)
        p.paragraph_format.right_indent = Inches(0)
        run = p.add_run(f"[Figure: {fig_path.name} — file not found]")
        _apply_font(run, italic=True, color=(180, 0, 0))
        return _add_caption(doc, caption, fig_label=fig_key)

    # Insert image — constrain to page text width (6.5" for 1" margins)
    para = doc.add_paragraph()
    _para_fmt(para, space_before=6, space_after=2,
              alignment=WD_ALIGN_PARAGRAPH.CENTER)
    para.paragraph_format.left_indent  = Inches(0)
    para.paragraph_format.right_indent = Inches(0)
    run  = para.add_run()
    run.add_picture(str(fig_path), width=Inches(min(width_in, 6.5)))
    cap = _add_caption(doc, caption, fig_label=fig_key)
    return cap


# ──────────────────────────────────────────────────────────────────────────────
# Markdown table → Word table
# ──────────────────────────────────────────────────────────────────────────────

def _parse_md_table(lines):
    """
    Parse a list of markdown table lines into (headers, rows).
    Separator lines (---) are skipped.
    """
    headers, rows = [], []
    for line in lines:
        if re.match(r'^\|[-| :]+\|$', line.strip()):
            continue  # separator
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if not headers:
            headers = cells
        else:
            rows.append(cells)
    return headers, rows


def _add_md_table(doc, md_lines, caption_text=None, is_simulated=False):
    """Convert markdown table lines to a Word table."""
    headers, rows = _parse_md_table(md_lines)
    if not headers:
        return
    n_cols = len(headers)
    n_rows = len(rows) + 1  # +1 for header

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    for j, h in enumerate(headers):
        cell = hdr_row.cells[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.clear()
        run  = para.add_run(h.replace("**", ""))
        _apply_font(run, size_pt=9, bold=True)
        _para_fmt(para, space_before=1, space_after=1,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
        # shade header
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "D9E1F2")
        tcPr.append(shd)

    # Data rows
    for i, row_data in enumerate(rows):
        tr = table.rows[i + 1]
        for j, cell_text in enumerate(row_data):
            if j >= n_cols:
                break
            cell = tr.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.clear()
            _render_inline(para, cell_text)
            _para_fmt(para, space_before=1, space_after=1,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT)
            for run in para.runs:
                _set_font_meta(run, size_pt=9)  # preserve bold/italic from inline

    # Table caption
    if caption_text:
        cap_para = doc.add_paragraph()
        _para_fmt(cap_para, space_before=2, space_after=8,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
        run = cap_para.add_run(caption_text)
        _apply_font(run, size_pt=9, italic=True)
        if is_simulated:
            _add_footnote(doc, cap_para, SIMULATED_FN_TEXT)


# ──────────────────────────────────────────────────────────────────────────────
# Section content extraction from markdown files
# ──────────────────────────────────────────────────────────────────────────────

def _extract_final(filepath, marker):
    """
    Extract text after the first occurrence of `marker` in `filepath`.
    Returns a list of stripped lines.
    """
    text  = Path(filepath).read_text(encoding="utf-8", errors="replace")
    idx   = text.find(marker)
    if idx == -1:
        print(f"  [WARN] Marker '{marker}' not found in {filepath.name}; using full file.")
        return text.splitlines()
    after = text[idx + len(marker):]
    # Strip any leading blank lines or dashes
    lines = after.splitlines()
    return lines


FINAL_MARKERS = {
    "abstract":     "## === FINAL ABSTRACT ===",
    "introduction": "## === FINAL INTRODUCTION ===",
    "related_work": "## FINAL RELATED WORK",
    "methodology":  "=== FINAL METHODOLOGY ===",
    "experiments":  "=== FINAL EXPERIMENTS ===",
    "analysis":     "# FINAL SECTIONS",
    "appendix":     "# Appendix",
    "disclaimer":   "## === FINAL DISCLAIMER ===",
    "references":   "## === FINAL REFERENCES ===",
}


def _get_section_lines(section_key):
    fpath   = SECTION_FILES[section_key]
    marker  = FINAL_MARKERS[section_key]
    lines   = _extract_final(fpath, marker)
    # Remove trailing revision-log / round notes / figures / bibtex sections
    clean   = []
    stop_markers = [
        "## REVISION LOG", "## Writing & Review", "=== ROUND", "## FIGURES",
        "## Complete BibTeX", "## Verification Notes", "## Additional Citations",
        "## BibTeX for", "## Citation Key",
    ]
    for line in lines:
        if any(line.strip().startswith(m) for m in stop_markers):
            break
        clean.append(line)
    return clean


# ──────────────────────────────────────────────────────────────────────────────
# Line-by-line markdown renderer into Word document
# ──────────────────────────────────────────────────────────────────────────────

def _render_lines(doc, lines, section_number_offset=0):
    """
    Render a list of markdown lines into the Word document.
    Handles: headings, paragraphs, bullet lists, numbered lists,
    blockquotes, horizontal rules, and markdown tables.
    """
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # --- Skip empty lines (handled via paragraph spacing)
        if not stripped:
            i += 1
            continue

        # --- Heading
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text  = stripped.lstrip("#").strip()
            # Remove markdown bold markers from headings
            text  = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            _add_heading(doc, text, level=min(level, 3))
            i += 1
            continue

        # --- Horizontal rule → thin paragraph spacing
        if stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            _para_fmt(p, space_before=4, space_after=4)
            r = p.add_run()
            r.add_break()
            i += 1
            continue

        # --- Markdown table
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            is_simulated = any("simulated" in tl.lower() or "*" in tl
                               for tl in table_lines)
            _add_md_table(doc, table_lines, is_simulated=is_simulated)
            continue

        # --- Blockquote (> text) → indented block; accumulate multi-line
        if stripped.startswith(">"):
            bq_parts = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                bq_parts.append(lines[i].strip().lstrip(">").strip())
                i += 1
            full_bq = " ".join(bq_parts).strip()
            # Skip reviewer/editor patch notes (not paper content)
            if re.match(r'\*{0,2}PATCH NOTE', full_bq):
                continue
            # Strip trailing patch-note attributions
            full_bq = re.sub(
                r'\s*--\s*(?:ACL|EMNLP|ICLR|NeurIPS|AAAI)\s+[Rr]ound\S*\s*\S*\s*$',
                '', full_bq
            ).strip()
            full_bq = re.sub(
                r'\s*--\s*[A-Za-z]+ round\d+ patch\s*\S*\s*$',
                '', full_bq
            ).strip()
            if not full_bq:
                continue
            para = doc.add_paragraph()
            _para_fmt(para, space_before=2, space_after=2,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
            para.paragraph_format.left_indent  = Inches(0.3)
            para.paragraph_format.right_indent = Inches(0.3)
            _render_inline(para, full_bq)
            for run in para.runs:
                _apply_font(run, size_pt=9,
                            bold=run.font.bold or False,
                            italic=not (run.font.bold or False))
            continue

        # --- Unordered bullet list
        if re.match(r'^[-*+]\s+', stripped):
            text = re.sub(r'^[-*+]\s+', '', stripped)
            para = doc.add_paragraph(style="List Bullet")
            _render_inline(para, text)
            for run in para.runs:
                _set_font_meta(run, size_pt=10)  # preserve bold/italic from inline
            _para_fmt(para, space_before=1, space_after=1,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
            i += 1
            continue

        # --- Numbered list
        if re.match(r'^\d+\.\s+', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            para = doc.add_paragraph(style="List Number")
            _render_inline(para, text)
            for run in para.runs:
                _set_font_meta(run, size_pt=10)  # preserve bold/italic from inline
            _para_fmt(para, space_before=1, space_after=1,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
            i += 1
            continue

        # --- Code block (```) → monospaced
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            for cl in code_lines:
                cp = doc.add_paragraph()
                _para_fmt(cp, space_before=0, space_after=0)
                cp.paragraph_format.left_indent = Inches(0.3)
                run = cp.add_run(cl)
                _apply_font(run, name="Courier New", size_pt=8)
            continue

        # --- Display math ($$...$$) → centered equation block
        if stripped.startswith("$$"):
            if stripped.endswith("$$") and len(stripped) > 4:
                # Single-line $$formula$$
                math_content = stripped[2:-2].strip()
                i += 1
            else:
                # Multi-line: collect until closing $$
                math_parts = []
                i += 1
                while i < len(lines):
                    ml = lines[i].strip()
                    if ml == "$$" or ml.startswith("$$"):
                        i += 1
                        break
                    math_parts.append(ml)
                    i += 1
                math_content = " ".join(math_parts)
            math_content = _preprocess_latex(math_content)
            para = doc.add_paragraph()
            _para_fmt(para, space_before=6, space_after=6,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
            run = para.add_run(math_content)
            _apply_font(run, size_pt=10, italic=True)
            continue

        # --- Figure placement tag: [FIG:fig_key]
        if re.match(r'^\[FIG:[^\]]+\]$', stripped):
            fig_key = re.match(r'^\[FIG:([^\]]+)\]$', stripped).group(1)
            _add_figure(doc, fig_key)
            i += 1
            continue

        # --- Normal paragraph
        # Collect continuation lines (non-empty, non-heading, non-list, etc.)
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_stripped = lines[i].strip()
            if (not next_stripped or
                    next_stripped.startswith("#") or
                    next_stripped.startswith("|") or
                    next_stripped.startswith(">") or
                    next_stripped.startswith("```") or
                    re.match(r'^[-*+]\s+', next_stripped) or
                    re.match(r'^\d+\.\s+', next_stripped) or
                    next_stripped in ("---", "***", "___")):
                break
            para_lines.append(next_stripped)
            i += 1
        full_text = " ".join(para_lines)
        is_simulated = "(simulated" in full_text.lower() or "simulated*" in full_text.lower()
        para = _add_normal(doc, full_text)
        if is_simulated:
            _add_footnote(doc, para, SIMULATED_FN_TEXT)

    return


# ──────────────────────────────────────────────────────────────────────────────
# Main document builder
# ──────────────────────────────────────────────────────────────────────────────

def build_document(output_path):
    print(f"[CoEval DocX Builder] Started: {datetime.now():%Y-%m-%d %H:%M:%S}")
    doc = Document()

    # ── Global default font (Times New Roman 10pt)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    nf = style._element.get_or_add_rPr().get_or_add_rFonts()
    nf.set(qn("w:ascii"),    "Times New Roman")
    nf.set(qn("w:hAnsi"),    "Times New Roman")
    nf.set(qn("w:eastAsia"), "Times New Roman")

    # ── Page setup: Letter, 1" margins, single column (title block)
    section0 = doc.sections[0]
    section0.page_width  = Inches(8.5)
    section0.page_height = Inches(11.0)
    _set_margins(section0)
    _set_cols(section0, num_cols=1)
    _add_page_numbers(section0)

    # ══════════════════════════════════════════════════════════════════════════
    # TITLE BLOCK (single-column)
    # ══════════════════════════════════════════════════════════════════════════
    print("  Building title block...")

    # Title — read from 00_paper_metadata.md ## TITLE
    _title_text = _META.get(
        "TITLE",
        "CoEval: A Self-Evaluating LLM Ensemble Framework for\n"
        "Scalable, Attribute-Controlled Benchmark Generation"
    )
    title_para = doc.add_paragraph()
    _para_fmt(title_para, space_before=0, space_after=6,
              alignment=WD_ALIGN_PARAGRAPH.CENTER)
    title_run = title_para.add_run(_title_text)
    _apply_font(title_run, size_pt=16, bold=True)

    # Authors line — read from ## AUTHORS
    _authors_text = _META.get("AUTHORS", "Alexander Apartsin\u00b9   Yehudit Aperstein\u00b2")
    authors_para = doc.add_paragraph()
    _para_fmt(authors_para, space_before=4, space_after=2,
              alignment=WD_ALIGN_PARAGRAPH.CENTER)
    auth_run = authors_para.add_run(_authors_text)
    _apply_font(auth_run, size_pt=11, bold=False)

    # Affiliations — read from ## AFFILIATION_1 and ## AFFILIATION_2
    _aff1_text = _META.get("AFFILIATION_1", "\u00b9 Holon Institute of Technology (HIT), Holon, Israel")
    _aff2_text = _META.get("AFFILIATION_2", "\u00b2 Afeka Tel Aviv Academic College of Engineering, Tel Aviv, Israel")
    aff1 = doc.add_paragraph()
    _para_fmt(aff1, space_before=0, space_after=0,
              alignment=WD_ALIGN_PARAGRAPH.CENTER)
    r1 = aff1.add_run(_aff1_text)
    _apply_font(r1, size_pt=9, italic=True)

    aff2 = doc.add_paragraph()
    _para_fmt(aff2, space_before=0, space_after=8,
              alignment=WD_ALIGN_PARAGRAPH.CENTER)
    r2 = aff2.add_run(_aff2_text)
    _apply_font(r2, size_pt=9, italic=True)

    # Code availability footnote — read from ## CODE_FOOTNOTE
    _code_fn_text = _META.get("CODE_FOOTNOTE", f"Code and data available at: {REPO_URL}")
    _add_footnote(doc, aff2, _code_fn_text)

    # ══════════════════════════════════════════════════════════════════════════
    # ABSTRACT (single-column, boxed-style via indented paragraph)
    # ══════════════════════════════════════════════════════════════════════════
    print("  Building abstract...")

    abs_label = doc.add_paragraph()
    _para_fmt(abs_label, space_before=6, space_after=2,
              alignment=WD_ALIGN_PARAGRAPH.CENTER)
    rl = abs_label.add_run("Abstract")
    _apply_font(rl, size_pt=11, bold=True)

    # Extract abstract text — stop at word-count notes, blockquotes, or HR
    abs_lines   = _get_section_lines("abstract")
    abs_text    = []
    for line in abs_lines:
        stripped = line.strip()
        if not stripped:
            continue  # skip blank separators; paragraphs are joined into one
        if stripped.startswith("#") or stripped.startswith(">") or stripped in ("---", "***", "___"):
            break
        if re.match(r'\*{1,2}Word\s*count', stripped, re.IGNORECASE):
            break
        abs_text.append(stripped)

    full_abstract = " ".join(abs_text)
    abs_para = doc.add_paragraph()
    _para_fmt(abs_para, space_before=0, space_after=8,
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    abs_para.paragraph_format.left_indent  = Inches(0.5)
    abs_para.paragraph_format.right_indent = Inches(0.5)
    _render_inline(abs_para, full_abstract)
    for run in abs_para.runs:
        _apply_font(run, size_pt=10)
    # Simulated footnote in abstract
    _add_footnote(doc, abs_para, SIMULATED_FN_TEXT)

    # Keywords line
    kw_para = doc.add_paragraph()
    _para_fmt(kw_para, space_before=2, space_after=12,
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    kw_para.paragraph_format.left_indent  = Inches(0.5)
    kw_para.paragraph_format.right_indent = Inches(0.5)
    _kw_text = _META.get(
        "KEYWORDS",
        "LLM evaluation, benchmark generation, ensemble scoring, "
        "LLM-as-judge, attribute-controlled generation, inter-rater agreement"
    )
    kw_run = kw_para.add_run(f"Keywords: {_kw_text}")
    _apply_font(kw_run, size_pt=9, italic=True)

    # Single-column continues throughout body (no column transition needed)

    # ══════════════════════════════════════════════════════════════════════════
    # §1 INTRODUCTION
    # ══════════════════════════════════════════════════════════════════════════
    print("  Building §1 Introduction...")
    _add_heading(doc, "1  Introduction", level=1)
    intro_lines = _get_section_lines("introduction")
    _render_lines(doc, intro_lines)

    # ══════════════════════════════════════════════════════════════════════════
    # Figure 1 — Architecture (placed after introduction)
    # ══════════════════════════════════════════════════════════════════════════
    _add_figure(doc, "fig1")

    # ══════════════════════════════════════════════════════════════════════════
    # §2 RELATED WORK
    # ══════════════════════════════════════════════════════════════════════════
    print("  Building §2 Related Work...")
    _add_heading(doc, "2  Related Work", level=1)
    rw_lines = _get_section_lines("related_work")
    _render_lines(doc, rw_lines)

    # ══════════════════════════════════════════════════════════════════════════
    # §3 METHODOLOGY
    # ══════════════════════════════════════════════════════════════════════════
    print("  Building §3 Methodology...")
    _add_heading(doc, "3  Framework and Methodology", level=1)
    meth_lines = _get_section_lines("methodology")
    _render_lines(doc, meth_lines)

    # Figure 2 — Teacher-Student-Judge diagram
    _add_figure(doc, "fig2_tsj")

    # ══════════════════════════════════════════════════════════════════════════
    # §4 EXPERIMENTS
    # ══════════════════════════════════════════════════════════════════════════
    print("  Building §4 Experiments...")
    _add_heading(doc, "4  Experiments and Results", level=1)
    exp_lines = _get_section_lines("experiments")
    _render_lines(doc, exp_lines)

    # Key result figures — representative screenshots and quantitative tables
    _add_figure(doc, "fig3_overview")      # overview dashboard (representative)
    _add_figure(doc, "fig4_judge_agree")   # judge agreement (key finding)
    _add_figure(doc, "fig11_kappa_matrix") # kappa matrix table
    _add_figure(doc, "fig12_teacher_disc") # teacher discrimination table
    _add_figure(doc, "fig13_cost")         # cost breakdown table

    # Simulated-result placeholder figures (marked with *)
    _add_figure(doc, "fig8_ensemble_abl")
    _add_figure(doc, "fig9_benchmark_comp")
    _add_figure(doc, "fig10_pos_bias")
    # Detailed report screenshots (fig5, fig6, fig7, fig14, fig15) moved to Appendix A.2

    # ══════════════════════════════════════════════════════════════════════════
    # §5 ANALYSIS & DISCUSSION / §6 LIMITATIONS / §7 CONCLUSION / ETHICS
    # ══════════════════════════════════════════════════════════════════════════
    print("  Building §5–§7 + Ethics...")
    anl_lines = _get_section_lines("analysis")
    _render_lines(doc, anl_lines)

    # ══════════════════════════════════════════════════════════════════════════
    # DISCLAIMER (end of paper, before references) — read from 10_disclaimer.md
    # ══════════════════════════════════════════════════════════════════════════
    print("  Adding disclaimer...")
    _add_heading(doc, "Disclaimer: Real vs. Simulated Results", level=1)
    disc_lines = _get_section_lines("disclaimer")
    _render_lines(doc, disc_lines)

    # ══════════════════════════════════════════════════════════════════════════
    # REFERENCES
    # ══════════════════════════════════════════════════════════════════════════
    print("  Building References...")
    _add_heading(doc, "References", level=1)
    ref_lines = _get_section_lines("references")
    # Render only the human-readable reference entries.
    # Skip: section headers (##, ###), VERIFIED blockquotes (>), BibTeX blocks,
    # horizontal rules, blank lines, and the file title line.
    in_bibtex = False
    for line in ref_lines:
        stripped = line.strip()
        if stripped.startswith("```bibtex") or stripped.startswith("```"):
            in_bibtex = stripped.startswith("```bibtex")
            continue
        if in_bibtex:
            continue
        if not stripped:
            continue
        if stripped.startswith("#") or stripped.startswith("---") or \
                stripped.startswith("*") and stripped.endswith("*") and len(stripped) < 6:
            continue
        if stripped.startswith(">"):
            continue   # skip VERIFIED notes and editorial comments
        # Reference entry — hanging indent, 9pt
        p = doc.add_paragraph()
        _para_fmt(p, space_before=2, space_after=2)
        p.paragraph_format.left_indent        = Inches(0.3)
        p.paragraph_format.first_line_indent  = Inches(-0.3)
        _render_inline(p, stripped)
        for run in p.runs:
            _set_font_meta(run, size_pt=9)  # preserve bold citation keys

    # ══════════════════════════════════════════════════════════════════════════
    # APPENDIX (single-column throughout — no section break needed)
    # ══════════════════════════════════════════════════════════════════════════
    print("  Building Appendix...")
    _add_heading(doc, "Appendix", level=1)
    app_lines = _get_section_lines("appendix")
    _render_lines(doc, app_lines)

    # A.1, A.2 sections (wizard + detailed reports) rendered via [FIG:key]
    # markers in 10_appendix.md — see _render_lines above.

    # ══════════════════════════════════════════════════════════════════════════
    # FOOTNOTES SECTION (collected superscripts)
    # ══════════════════════════════════════════════════════════════════════════
    print("  Writing footnotes section...")
    _add_heading(doc, "Footnotes", level=2)
    _FN.write_to_doc(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # Save
    # ══════════════════════════════════════════════════════════════════════════
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    size_kb = output_path.stat().st_size // 1024
    print(f"\n[CoEval DocX Builder] Done: {output_path}")
    print(f"  File size: {size_kb} KB")
    print(f"  Figures inserted: {sum(1 for k in FIGURES if Path(FIGURES[k][0]).exists())}/{len(FIGURES)}")


# ──────────────────────────────────────────────────────────────────────────────
# Entry point
# ──────────────────────────────────────────────────────────────────────────────

def _next_version(drafts_dir: Path) -> str:
    """
    Auto-detect the highest existing vX.Y version in drafts_dir and return vX.(Y+1).
    Falls back to v1.0 if no versioned files exist.
    """
    import re as _re
    pat = _re.compile(r'CoEval_ACL2026_v(\d+)\.(\d+)\.docx$')
    max_major, max_minor = 1, -1
    for f in drafts_dir.glob("CoEval_ACL2026_v*.docx"):
        m = pat.match(f.name)
        if m:
            maj, min_ = int(m.group(1)), int(m.group(2))
            if (maj, min_) > (max_major, max_minor):
                max_major, max_minor = maj, min_
    if max_minor == -1:
        return "v1.0"
    return f"v{max_major}.{max_minor + 1}"


if __name__ == "__main__":
    DRAFTS_DIR = PAPER_DIR / "drafts"
    DRAFTS_DIR.mkdir(parents=True, exist_ok=True)

    next_ver = _next_version(DRAFTS_DIR)
    default_out = str(DRAFTS_DIR / f"CoEval_ACL2026_{next_ver}.docx")

    parser = argparse.ArgumentParser(description="Generate CoEval paper docx")
    parser.add_argument(
        "--output", "-o",
        default=default_out,
        help=f"Output .docx path (default: auto-versioned → {default_out})"
    )
    args = parser.parse_args()
    build_document(args.output)
